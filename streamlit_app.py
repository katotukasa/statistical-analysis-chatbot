import streamlit as st
import google.generativeai as genai
import os
from pypdf import PdfReader 
import pandas as pd
import matplotlib.pyplot as plt
from matplotlib import font_manager
from docx import Document
from docx.shared import Inches
from io import BytesIO

# ==========================================================
# Matplotlibで日本語フォントを設定
# ==========================================================
FONT_PATH = 'ipaexg.ttf' 
try:
    font_prop = font_manager.FontProperties(fname=FONT_PATH)
    plt.rcParams['font.family'] = font_prop.get_name()
    plt.rcParams['axes.unicode_minus'] = False 
    st.session_state.font_prop = font_prop
except FileNotFoundError:
    st.warning(f"警告: 日本語フォントファイル '{FONT_PATH}' が見つかりません。グラフの日本語が文字化けする可能性があります。")
    plt.rcParams['font.family'] = 'DejaVu Sans'
    st.session_state.font_prop = None
# ==========================================================


# --- アプリケーションの基本設定 ---
st.set_page_config(
    page_title="統計分析支援チャットボット",
    page_icon="🤖",
    layout="wide",
)

st.title("📊 統計分析支援チャットボット")
st.write(
    "ようこそ！このチャットボットは、あなたがアップロードした文書（統計分析の計画など）やデータ（CSVファイル）に基づいて、**記述統計**、**グラフ化**、**推奨統計処理**の提案を行います。"
)
st.write(
    "まずは、お持ちのGemini APIキーを入力し、分析計画が書かれたファイルまたはCSVファイルをアップロードしてください。"
)

# --- サイドバーでのAPIキー入力 ---
with st.sidebar:
    gemini_api_key = st.text_input("Gemini API Key", type="password", key="gemini_api_key")
    "[Gemini APIキーを取得する](https://aistudio.google.com/app/apikey)"

# --- メインコンテンツ ---
if not gemini_api_key:
    st.info("サイドバーからGemini APIキーを入力してください。")
    st.stop()

# APIキーの認証
try:
    genai.configure(api_key=gemini_api_key)
    model = genai.GenerativeModel("gemini-2.5-flash") 
except Exception as e:
    st.error(f"APIキーの認証に失敗しました。正しいキーを入力してください。: {e}")
    st.stop()

# --- プロンプト設定 ---
SYSTEM_PROMPT = """
あなたは、統計分析の専門家であり、教育者です。
ユーザーから提供された文書（研究計画、分析のメモ、データ構造の概要など）を深く理解し、以下の役割を担ってください。

1.  **記述統計とグラフの解説**: 提供されたCSVファイルの記述統計結果やグラフの内容を、分析の文脈に沿って分かりやすく解説します。
2.  **推奨統計処理の提案**: ドキュメントの内容とデータの特性（記述統計、グラフ）に基づき、最も適切だと思われる統計手法を複数提案し、それぞれのメリット・デメリットを分かりやすく説明します。
3.  **質問応答**: 統計学の概念、特定の手法、ツールの使い方（例：Pythonのライブラリ）など、ユーザーからのあらゆる質問に、初心者にも理解できるように丁寧に答えます。
4.  **対話の記憶**: 過去の会話を記憶し、文脈に沿った対話を続けます。

あなたの目的は、ユーザーが自身の研究や学習において、統計分析を正しく、かつ自信を持って活用できるようになることを支援することです。
"""

# --- PDFファイルからテキストを抽出する関数 ---
def read_pdf_text(pdf_file):
    """
    アップロードされたPDFファイルからすべてのページテキストを抽出する
    """
    try:
        reader = PdfReader(pdf_file)
        text = ""
        for page in reader.pages:
            extracted_text = page.extract_text()
            if extracted_text:
                text += extracted_text + "\n\n"
        return text
    except Exception as e:
        st.error(f"PDFの読み込み中にエラーが発生しました: {e}")
        return ""

# --- CSVファイルから構造と記述統計を抽出する関数 ---
def get_csv_analysis_text(csv_file):
    """
    アップロードされたCSVファイルから構造、記述統計を抽出し、データフレームをセッションに保存する
    """
    try:
        csv_file.seek(0)
        df = pd.read_csv(csv_file)
        st.session_state.data_df = df
        
        col_info = "\n".join([f"- {col}: {dtype}" for col, dtype in df.dtypes.items()])
        desc_stats = df.describe(include='all').to_markdown()

        content = (
            f"これは、アップロードされたCSVファイル「{csv_file.name}」のデータ構造と記述統計の概要です。\n"
            f"行数: {len(df)}、列数: {len(df.columns)}\n\n"
            f"### カラム名とデータ型:\n{col_info}\n\n"
            f"### 記述統計の結果:\n{desc_stats}"
        )
        return content
    except Exception as e:
        st.error(f"CSVファイルの読み込みまたは処理中にエラーが発生しました: {e}")
        st.session_state.data_df = pd.DataFrame()
        return ""

# --- グラフ生成およびセッションへの保存機能 (st.pyplotは呼ばない) ---
def generate_and_store_plots(df):
    """
    グラフを生成し、図オブジェクトと画像バッファをセッションに保存する。
    利用可能なグラフのタイトルリストを返す。
    """
    st.session_state.plot_images = {} # Wordレポート用バッファ
    st.session_state.plot_figures = {} # Streamlit表示用figureオブジェクト
    available_plots = []
    
    font_prop = st.session_state.get('font_prop', None)

    numeric_cols = df.select_dtypes(include=['number']).columns.tolist()
    object_cols = df.select_dtypes(include=['object']).columns.tolist()
    
    # 1. 数値型データのヒストグラム/箱ひげ図
    if numeric_cols:
        for col in numeric_cols[:4]: # サンプルとして最初の4つ
            # ヒストグラム
            hist_title = f'{col} のヒストグラム'
            fig_hist, ax_hist = plt.subplots(figsize=(4, 3))
            ax_hist.hist(df[col].dropna(), bins='auto', edgecolor='black')
            ax_hist.set_title(hist_title, fontproperties=font_prop if font_prop else None)
            
            # 保存
            hist_buf = BytesIO()
            fig_hist.savefig(hist_buf, format='png')
            st.session_state.plot_images[hist_title] = hist_buf
            st.session_state.plot_figures[hist_title] = fig_hist
            available_plots.append(hist_title)
            
            # 箱ひげ図
            box_title = f'{col} の箱ひげ図'
            fig_box, ax_box = plt.subplots(figsize=(4, 3))
            ax_box.boxplot(df[col].dropna())
            ax_box.set_title(box_title, fontproperties=font_prop if font_prop else None)
            
            # 保存
            box_buf = BytesIO()
            fig_box.savefig(box_buf, format='png')
            st.session_state.plot_images[box_title] = box_buf
            st.session_state.plot_figures[box_title] = fig_box
            available_plots.append(box_title)

    # 2. カテゴリ型データの度数分布
    if object_cols:
        for col in object_cols[:2]: # サンプルとして最初の2つ
            bar_title = f'{col} の度数分布'
            counts = df[col].value_counts().head(10)
            fig_bar, ax_bar = plt.subplots(figsize=(6, 4))
            ax_bar.bar(counts.index.astype(str), counts.values)
            ax_bar.set_title(bar_title, fontproperties=font_prop if font_prop else None)
            
            # フォント設定
            if font_prop:
                for label in ax_bar.get_xticklabels():
                    label.set_fontproperties(font_prop)
                for label in ax_bar.get_yticklabels():
                    label.set_fontproperties(font_prop)
            
            ax_bar.tick_params(axis='x', rotation=45)
            plt.tight_layout()
            
            # 保存
            bar_buf = BytesIO()
            fig_bar.savefig(bar_buf, format='png')
            st.session_state.plot_images[bar_title] = bar_buf
            st.session_state.plot_figures[bar_title] = fig_bar
            available_plots.append(bar_title)

    st.session_state.available_plots = available_plots
    
    # 最後に、生成したfigureオブジェクトを全てクローズしてメモリを解放（st.session_stateに保存されているため問題ない）
    for fig in st.session_state.plot_figures.values():
        plt.close(fig)

# --- Wordレポート生成関数 ---
def create_word_report(analysis_content, summary_content, plot_images):
    """
    AIの提案と記述統計、グラフをWordファイルとして生成する
    """
    document = Document()
    document.add_heading('統計分析レポート', 0)
    document.add_paragraph(f'作成日時: {pd.Timestamp.now().strftime("%Y年%m月%d日 %H:%M:%S")}')
    document.add_paragraph('---')

    # 1. AIによる推奨統計処理の提案
    document.add_heading('1. AIによる推奨統計処理の提案', level=1)
    
    for line in summary_content.split('\n'):
        if line.startswith('#'):
            level = line.count('#')
            if level <= 3:
                document.add_heading(line.lstrip('# ').strip(), level=level + 1)
        elif line.strip():
            document.add_paragraph(line)
            
    document.add_paragraph('---')

    # 2. アップロードされたファイルの概要/記述統計
    document.add_heading('2. ファイル概要と記述統計', level=1)
    
    document.add_paragraph(analysis_content)
    document.add_paragraph('---')

    # 3. グラフ
    if plot_images:
        document.add_heading('3. データのグラフ', level=1)
        from docx.shared import Inches 
        for key, buf in plot_images.items():
            document.add_heading(key.replace('_', ' ').title(), level=2)
            buf.seek(0)
            document.add_picture(buf, width=Inches(3.0)) 
    
    # WordファイルをBytesIOストリームに保存
    doc_io = BytesIO()
    document.save(doc_io)
    doc_io.seek(0)
    return doc_io.getvalue()


# ファイルアップローダー
uploaded_file = st.file_uploader(
    "分析計画のファイル（.md、.txt、.pdf）またはデータファイル（.csv）をアップロードしてください",
    type=["md", "txt", "pdf", "csv"]
)

# 1. ファイルアップロード/変更の処理
if uploaded_file is not None:
    if "last_uploaded_filename" not in st.session_state or st.session_state.last_uploaded_filename != uploaded_file.name:
        
        st.session_state.last_uploaded_filename = uploaded_file.name
        st.session_state.messages = []
        st.session_state.summary = None 
        st.session_state.data_df = pd.DataFrame()
        st.session_state.plot_images = {}
        st.session_state.plot_figures = {} # 新しいfigureオブジェクトをリセット
        st.session_state.available_plots = [] # 利用可能なグラフリストをリセット

        file_extension = uploaded_file.name.split(".")[-1].lower()
        st.session_state.document_content = ""
        
        try:
            with st.spinner(f"「{uploaded_file.name}」の内容を読み込み中..."):
                if file_extension in ["md", "txt"]:
                    uploaded_file.seek(0)
                    st.session_state.document_content = uploaded_file.read().decode("utf-8")
                elif file_extension == "pdf":
                    st.session_state.document_content = read_pdf_text(uploaded_file)
                elif file_extension == "csv":
                    # CSVの場合は、記述統計結果とデータフレームをセッションに格納
                    st.session_state.document_content = get_csv_analysis_text(uploaded_file)
                else:
                    st.error("サポートされていないファイル形式です。")
                    st.session_state.document_content = ""
                    st.stop()
            
            if st.session_state.document_content:
                st.success(f"「{uploaded_file.name}」の読み込みが完了しました。")
            else:
                st.warning(f"「{uploaded_file.name}」から内容を抽出できませんでした。ファイル内容を確認してください。")
                st.session_state.document_content = ""
                st.stop()

        except Exception as e:
            st.error(f"ファイル内容の処理中に致命的なエラーが発生しました: {e}")
            st.session_state.document_content = ""
            st.stop()
    
# 2. メインロジック（レジューム処理）
if "document_content" in st.session_state and st.session_state.document_content:

    uploaded_file_name = st.session_state.last_uploaded_filename
    is_csv_file = uploaded_file_name.split(".")[-1].lower() == "csv"

    # --- 記述統計とグラフの準備/表示 ---
    if is_csv_file and not st.session_state.data_df.empty:
        with st.expander("📚 CSVデータ構造と記述統計の結果", expanded=True):
            st.markdown(st.session_state.document_content)
            
        # グラフデータがセッションにまだない場合（初回ロード時）のみグラフを生成
        if not st.session_state.available_plots:
            with st.spinner("グラフを生成しています..."):
                generate_and_store_plots(st.session_state.data_df)

        st.subheader("📊 データのグラフ化")
        
        # グラフ選択リスト
        selected_plots = st.multiselect(
            "表示するグラフを選択してください",
            st.session_state.available_plots,
            default=st.session_state.available_plots # デフォルトで全て選択
        )
        
        # 選択されたグラフの表示
        if selected_plots:
            # グラフを2列で表示
            cols = st.columns(2)
            col_index = 0
            for plot_title in selected_plots:
                with cols[col_index % 2]:
                    fig = st.session_state.plot_figures.get(plot_title)
                    if fig:
                        # figを再利用して表示
                        st.pyplot(fig)
                col_index += 1
        else:
            st.info("表示するグラフを選択してください。")
        
    # --- AIによる推奨処理の提案 ---
    if st.session_state.document_content and not st.session_state.summary:
        with st.spinner("AIが推奨統計処理の提案を作成しています..."):
            try:
                if is_csv_file:
                    summary_prompt = (
                        f"{SYSTEM_PROMPT}\n\n---\n\n"
                        f"以下のCSVデータ構造の概要と記述統計の結果に基づき、このデータで可能な**推奨統計処理を簡潔に提案し、そのメリットを説明してください。**\n\n"
                        f"{st.session_state.document_content}"
                    )
                    expander_title = "アップロードされたデータに基づいた推奨統計処理の提案"
                else:
                    summary_prompt = f"{SYSTEM_PROMPT}\n\n---\n\n以下のドキュメントを3〜5行で簡潔に要約し、文書内容に基づいた統計手法の候補を提案してください。\n\n{st.session_state.document_content}"
                    expander_title = "アップロードされたドキュメントの要約と統計手法の候補"

                response = model.generate_content(summary_prompt)
                st.session_state.summary = response.text
                st.session_state.expander_title = expander_title

            except Exception as e:
                st.error(f"AIによる提案の生成中にエラーが発生しました: {e}")

    # --- AI提案の表示とWordダウンロードボタン ---
    if st.session_state.summary:
        with st.expander(st.session_state.expander_title, expanded=True):
            st.markdown(st.session_state.summary)

        # Wordレポートの生成とダウンロード
        if st.session_state.document_content and st.session_state.summary:
            report_data = create_word_report(
                st.session_state.document_content, 
                st.session_state.summary, 
                st.session_state.get('plot_images', {})
            )
            
            base_name = os.path.splitext(st.session_state.last_uploaded_filename)[0]
            download_file_name = f"{base_name}_分析レポート.docx"
            
            st.download_button(
                label="📄 Wordレポート (.docx) をダウンロード",
                data=report_data,
                file_name=download_file_name,
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

    # --- チャット機能 ---
    if "messages" not in st.session_state:
        st.session_state.messages = []

    # 過去のメッセージを表示
    for message in st.session_state.messages:
        with st.chat_message(message["role"]):
            st.markdown(message["content"])

    # ユーザーからの新しいメッセージ
    if prompt := st.chat_input("ファイル内容や分析結果について質問してください"):
        st.session_state.messages.append({"role": "user", "content": prompt})
        with st.chat_message("user"):
            st.markdown(prompt)

        try:
            with st.chat_message("assistant"):
                with st.spinner("AIが応答を生成中です..."):
                    full_prompt = (
                        f"{SYSTEM_PROMPT}\n\n"
                        f"--- 以下はアップロードされたファイルの内容（またはCSVの記述統計）です ---\n"
                        f"{st.session_state.document_content}\n\n"
                        f"--- 以下はこれまでの会話履歴と現在の質問です ---\n"
                    )
                    for msg in st.session_state.messages:
                        full_prompt += f"{msg['role']}: {msg['content']}\n"

                    response_stream = model.generate_content(full_prompt, stream=True)
                    
                    full_response = ""
                    response_placeholder = st.empty()
                    for chunk in response_stream:
                        if chunk.text:
                            full_response += chunk.text
                            response_placeholder.markdown(full_response + " ▌")
                    response_placeholder.markdown(full_response)

            st.session_state.messages.append({"role": "assistant", "content": full_response})

        except Exception as e:
            st.error(f"応答の生成中にエラーが発生しました: {e}")

else:
    # document_content がない場合（初回ロード時など）
    st.info("ファイルをアップロードすると、チャットが開始できます。")
